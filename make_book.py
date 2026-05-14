#!/usr/bin/env python3
"""
Genera Historia en 20 Minutos.docx combinando los Blog.md de episodios 001–020.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Configuración ──────────────────────────────────────────────────────────────

BOOK_TITLE    = "Historia en 20 Minutos"
FONT          = "Liberation Serif"
BODY_PT       = 12
HEADING_PT    = 14
CHAP_TITLE_PT = 22
EXCERPT_PT    = 13
HOOK_PT       = 12
CHAP_NUM_PT   = 36
HEADER_PT     = 10
PAGE_NUM_PT   = 10

BASE = Path("/Users/andresaguilar/source/podcast-en20minutos/Historia")

EPISODE_DIRS = [
    "001 Revolucion Francesa",
    "002 Imperio austrohungaro",
    "003 Servicio Postal",
    "004 Invasiones Inglesas",
    "005 Vikingos",
    "006 Canal de Panama",
    "007 El Renacimiento en Florencia",
    "008 - Alejangro Magno",
    "009 - La guerra de Vietnam",
    "010 - La caida de Constantinopla",
    "011 - Los Borgia",
    "012 - El Cumpleaños",
    "013 - El Feudalismo",
    "014 - La revolucion industrial",
    "015 - El Imperio Otomano",
    "016 - La independencia de Estados Unidos",
    "017 - La unificacion de Italia",
    "018 - El Imperio Persa",
    "019 - El Antiguo Egipto",
    "020 - El Imperio Inca",
]

PROLOGUE = """\
Este libro reúne veinte capítulos escritos a partir de una premisa simple: que la historia, \
bien contada, no necesita más de veinte minutos para capturar la atención de cualquier persona. \
Sin gráficos innecesarios, sin tecnicismos, sin la solemnidad que a veces intimida. \
Solo el relato de lo que pasó, por qué importó entonces, y por qué sigue importando hoy.

Cada capítulo fue concebido como una unidad autónoma. Puede leerse en cualquier orden, \
en cualquier momento. La Revolución Francesa no requiere haber leído sobre los Vikingos. \
El Imperio Inca no necesita contexto previo sobre el Imperio Persa. Pero quien los lea todos \
descubrirá que hay hilos invisibles que conectan estas historias: el mundo es una red de causas \
y consecuencias que no respeta fronteras ni siglos.

Este libro está pensado para el lector que quiere entender el mundo sin necesitar una carrera \
universitaria en Historia. Para quien siente curiosidad pero dispone de poco tiempo. \
Para quien alguna vez se preguntó por qué las cosas son como son y sospechó que la respuesta \
estaba en algún punto del pasado.

La historia no es una colección de fechas y batallas. Es el registro de cómo los seres humanos \
resolvieron —o no supieron resolver— los problemas que se les presentaron. Eso no cambió. \
Cambiaron los problemas, cambiaron las herramientas, pero la naturaleza humana que los enfrenta \
es reconociblemente la misma de siempre. Por eso la historia sigue siendo relevante. \
Por eso vale la pena contarla."""


# ── Parser de Markdown ─────────────────────────────────────────────────────────

def parse_blog(filepath):
    """
    Parsea Blog.md y devuelve:
      title:    str
      excerpt:  str  (de **Excerpt:** o del primer párrafo itálico)
      elements: list[(tipo, texto)]  — tipos: heading | italic | body | sep
    """
    text = Path(filepath).read_text(encoding="utf-8")
    lines = text.split("\n")

    # Saltar líneas de metadatos generadas por IA antes del contenido real
    start = 0
    for i, line in enumerate(lines):
        if re.match(r"^#+\s", line) or re.match(r"^\*\*Excerpt:\*\*", line):
            start = i
            break
    lines = lines[start:]

    # Extraer excerpt explícito (episodios 013-020)
    excerpt = None
    if lines and re.match(r"^\*\*Excerpt:\*\*", lines[0].strip()):
        m = re.match(r"^\*\*Excerpt:\*\*\s*(.+)", lines[0])
        if m:
            excerpt = m.group(1).strip()
        lines = lines[1:]

    # Saltar líneas en blanco
    while lines and not lines[0].strip():
        lines.pop(0)

    # Extraer título del primer encabezado #
    title = None
    if lines and re.match(r"^#+\s", lines[0]):
        raw = re.sub(r"^#+\s+", "", lines[0])
        title = re.sub(r"\*+", "", raw).strip()
        lines = lines[1:]

    # Parsear elementos
    elements = []
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        # Encabezado de sección (## o más profundo)
        if re.match(r"^#+\s", s):
            heading = re.sub(r"^#+\s+", "", s)
            heading = re.sub(r"\*+", "", heading).strip()
            elements.append(("heading", heading))
            i += 1
            continue

        # Separador ---
        if s == "---":
            elements.append(("sep", ""))
            i += 1
            continue

        # Párrafo itálico: empieza con * pero no con **
        if s.startswith("*") and not s.startswith("**"):
            block = []
            while i < len(lines) and lines[i].strip():
                block.append(lines[i].strip())
                i += 1
            joined = " ".join(block)
            joined = re.sub(r"^\*+", "", joined)
            joined = re.sub(r"\*+$", "", joined)
            elements.append(("italic", joined.strip()))
            continue

        # Párrafo de cuerpo
        block = []
        while i < len(lines):
            l = lines[i].strip()
            if not l:
                break
            if re.match(r"^#+\s", l) or l == "---":
                break
            if l.startswith("*") and not l.startswith("**"):
                break
            block.append(l)
            i += 1
        if block:
            elements.append(("body", " ".join(block)))

    # Si no hay excerpt explícito, usar el primer elemento itálico
    if not excerpt:
        for t, txt in elements:
            if t == "italic":
                excerpt = txt
                break  # El elemento permanece en la lista (hook + excerpt)

    return {"title": title, "excerpt": excerpt, "elements": elements}


# ── Helpers de DOCX ───────────────────────────────────────────────────────────

def page_num_field(para):
    """Inserta el campo PAGE en un párrafo de footer."""
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.font.name = FONT
    run.font.size = Pt(PAGE_NUM_PT)
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t"); placeholder.text = "0"
    fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "end")
    run._r.extend([fld1, instr, fld2, placeholder, fld3])


def set_hf_text(hf, text, italic=False):
    """Escribe texto en header/footer."""
    p = hf.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(HEADER_PT)
    run.italic = italic


def blank_hf(hf):
    """Vacía un header/footer."""
    hf.paragraphs[0].clear()


def enable_even_odd(doc):
    """Habilita headers distintos para páginas pares e impares a nivel de documento."""
    doc.settings.element.append(OxmlElement("w:evenAndOddHeaders"))


def set_margins(sec, left=3.0, right=2.5, top=2.5, bottom=2.5):
    sec.left_margin     = Cm(left)
    sec.right_margin    = Cm(right)
    sec.top_margin      = Cm(top)
    sec.bottom_margin   = Cm(bottom)
    sec.header_distance = Cm(1.25)
    sec.footer_distance = Cm(1.25)


def setup_chapter_section(sec, chapter_name):
    """
    Configura la sección de un capítulo:
    - Primera página (portada del cap.): sin header, con folio
    - Páginas impares: nombre del capítulo (itálica)
    - Páginas pares: título del libro
    - Footer en todas las páginas: número de página
    """
    sec.different_first_page_header_footer = True

    # Primera página — sin header
    sec.first_page_header.is_linked_to_previous = False
    blank_hf(sec.first_page_header)

    # Primera página — folio en footer
    sec.first_page_footer.is_linked_to_previous = False
    page_num_field(sec.first_page_footer.paragraphs[0])

    # Páginas impares — nombre del capítulo
    sec.header.is_linked_to_previous = False
    set_hf_text(sec.header, chapter_name, italic=True)

    # Páginas pares — título del libro
    sec.even_page_header.is_linked_to_previous = False
    set_hf_text(sec.even_page_header, BOOK_TITLE)

    # Footer impares
    sec.footer.is_linked_to_previous = False
    page_num_field(sec.footer.paragraphs[0])

    # Footer pares
    sec.even_page_footer.is_linked_to_previous = False
    page_num_field(sec.even_page_footer.paragraphs[0])


# ── Constructor de párrafos ───────────────────────────────────────────────────

def add_para(doc, text, *, pt=BODY_PT, bold=False, all_italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=Pt(0), space_after=Pt(6),
             left_indent=None, right_indent=None, first_line=None):
    """
    Agrega un párrafo con soporte de markup inline **negrita** y *itálica*.
    `all_italic=True` hace itálico todo el párrafo (más **negrita** interna).
    """
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after  = space_after
    if left_indent  is not None: pf.left_indent       = left_indent
    if right_indent is not None: pf.right_indent      = right_indent
    if first_line   is not None: pf.first_line_indent = first_line

    # Tokenizer inline: **bold**, *italic*, texto plano
    for m in re.finditer(r"(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)", text):
        chunk = m.group()
        run = p.add_run()
        run.font.name = FONT
        run.font.size = Pt(pt)
        if chunk.startswith("**") and chunk.endswith("**"):
            run.text   = chunk[2:-2]
            run.bold   = True
            run.italic = all_italic
        elif chunk.startswith("*") and chunk.endswith("*") and not all_italic:
            run.text   = chunk[1:-1]
            run.italic = True
            run.bold   = bold
        else:
            run.text   = chunk
            run.bold   = bold
            run.italic = all_italic

    # Fallback si el texto no tiene marcadores
    if not p.runs:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(pt)
        run.bold = bold
        run.italic = all_italic

    return p


# ── Renderizador de capítulos ─────────────────────────────────────────────────

def render_chapter(doc, data, num):
    """Renderiza la portada del capítulo + contenido en el documento."""
    title   = data["title"] or f"Capítulo {num}"
    excerpt = data.get("excerpt") or ""
    elems   = data["elements"]

    # ── Portada del capítulo (primera página de la sección, sin header) ──────

    # Espaciado vertical inicial
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(108)
    sp.paragraph_format.space_after  = Pt(0)

    # Número de capítulo en gris
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"{num:02d}")
    run.font.name      = FONT
    run.font.size      = Pt(CHAP_NUM_PT)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    # Título del capítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(24)
    run = p.add_run(title)
    run.font.name = FONT
    run.font.size = Pt(CHAP_TITLE_PT)
    run.bold      = True

    # Excerpt
    if excerpt:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(2.0)
        p.paragraph_format.right_indent = Cm(2.0)
        run = p.add_run(excerpt)
        run.font.name = FONT
        run.font.size = Pt(EXCERPT_PT)
        run.italic    = True

    # Salto de página (mismo section → las siguientes páginas ya tienen header)
    doc.add_page_break()

    # ── Contenido del capítulo ────────────────────────────────────────────────

    first_elem = True
    prev_type  = None

    for etype, etext in elems:

        if etype == "sep":
            prev_type = "sep"
            continue  # El espaciado entre párrafos maneja la separación visual

        if etype == "heading":
            add_para(doc, etext,
                     pt=HEADING_PT, bold=True,
                     align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=Pt(20), space_after=Pt(8))
            prev_type = "heading"

        elif etype == "italic":
            if first_elem:
                # Hook (primer párrafo itálico): ancho completo, justificado
                add_para(doc, etext,
                         pt=HOOK_PT, all_italic=True,
                         align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=Pt(12), space_after=Pt(16))
            else:
                # Pull quote: centrado, con sangría lateral
                add_para(doc, etext,
                         pt=BODY_PT, all_italic=True,
                         align=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(20), space_after=Pt(20),
                         left_indent=Cm(1.5), right_indent=Cm(1.5))
            prev_type = "italic"

        elif etype == "body":
            # Primera línea indentada solo si el párrafo anterior también era cuerpo
            fi = Cm(0.75) if prev_type == "body" else None
            add_para(doc, etext,
                     pt=BODY_PT,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     space_before=Pt(0), space_after=Pt(6),
                     first_line=fi)
            prev_type = "body"

        first_elem = False


# ── Función principal ─────────────────────────────────────────────────────────

def make_book():
    doc = Document()
    enable_even_odd(doc)

    # Estilo Normal base
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(BODY_PT)

    # ── SECCIÓN 0: Portada del libro ──────────────────────────────────────────

    s0 = doc.sections[0]
    set_margins(s0)
    s0.different_first_page_header_footer = True

    # Sin header ni footer en la portada
    s0.first_page_header.is_linked_to_previous = False
    blank_hf(s0.first_page_header)
    s0.first_page_footer.is_linked_to_previous = False
    blank_hf(s0.first_page_footer)

    # Contenido de la portada
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(144)
    sp.paragraph_format.space_after  = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run("En 20 Minutos")
    run.font.name = FONT; run.font.size = Pt(14); run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)
    run = p.add_run("Historia")
    run.font.name = FONT; run.font.size = Pt(42); run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run("Veinte capítulos para entender el mundo")
    run.font.name = FONT; run.font.size = Pt(13); run.italic = True

    # ── SECCIÓN 1: Prólogo ────────────────────────────────────────────────────

    s_pro = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(s_pro)
    s_pro.different_first_page_header_footer = True

    for hf in (s_pro.first_page_header, s_pro.header, s_pro.even_page_header):
        hf.is_linked_to_previous = False
        blank_hf(hf)

    for ftr in (s_pro.first_page_footer, s_pro.footer, s_pro.even_page_footer):
        ftr.is_linked_to_previous = False
        page_num_field(ftr.paragraphs[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after  = Pt(28)
    run = p.add_run("Prólogo")
    run.font.name = FONT; run.font.size = Pt(HEADING_PT); run.bold = True

    for blk in PROLOGUE.split("\n\n"):
        add_para(doc, blk.strip(),
                 space_before=Pt(0), space_after=Pt(10),
                 first_line=Cm(0.75))

    # ── Parsear todos los capítulos (necesario antes de construir el índice) ──

    print("Parseando episodios...", flush=True)
    chapters = []
    for ep in EPISODE_DIRS:
        path = BASE / ep / "Blog.md"
        ch = parse_blog(path)
        chapters.append(ch)
        print(f"  ✓ {ep[:30]:<30}  →  {ch['title'][:50] if ch['title'] else 'sin título'}")

    # ── SECCIÓN 2: Índice ─────────────────────────────────────────────────────

    s_toc = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(s_toc)
    s_toc.different_first_page_header_footer = True

    for hf in (s_toc.first_page_header, s_toc.header, s_toc.even_page_header):
        hf.is_linked_to_previous = False
        blank_hf(hf)

    for ftr in (s_toc.first_page_footer, s_toc.footer, s_toc.even_page_footer):
        ftr.is_linked_to_previous = False
        page_num_field(ftr.paragraphs[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after  = Pt(28)
    run = p.add_run("Índice")
    run.font.name = FONT; run.font.size = Pt(HEADING_PT); run.bold = True

    for idx, ch in enumerate(chapters, 1):
        t = ch["title"] or f"Capítulo {idx}"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.left_indent  = Cm(0.5)
        run = p.add_run(f"{idx:02d}.  ")
        run.font.name = FONT; run.font.size = Pt(BODY_PT); run.bold = True
        run = p.add_run(t)
        run.font.name = FONT; run.font.size = Pt(BODY_PT)

    # ── SECCIONES 3–22: Capítulos ─────────────────────────────────────────────

    print("\nGenerando capítulos...", flush=True)
    for idx, ch in enumerate(chapters, 1):
        name = ch["title"] or f"Capítulo {idx}"
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        set_margins(sec)
        setup_chapter_section(sec, name)
        render_chapter(doc, ch, idx)
        print(f"  ✓ Capítulo {idx:02d}: {name[:55]}")

    # ── Guardar ───────────────────────────────────────────────────────────────

    out = BASE.parent / "Historia en 20 Minutos.docx"
    doc.save(out)
    print(f"\n✓ Libro guardado en: {out}")
    return out


if __name__ == "__main__":
    make_book()
